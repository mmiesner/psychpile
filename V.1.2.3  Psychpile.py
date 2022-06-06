import PySimpleGUI as sg
from docx import Document
from docx.shared import Pt
import os
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
import zipfile
import shutil

# Create a custom logger
logger = logging.getLogger(__name__)

# Create handlers
c_handler = logging.StreamHandler()
f_handler = logging.FileHandler('file.log')
c_handler.setLevel(logging.WARNING)
f_handler.setLevel(logging.ERROR)

# Create formatters and add it to handlers
c_format = logging.Formatter('%(name)s - %(levelname)s - %(message)s')
f_format = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
#end logging module;start functions

def create_folder():
    os.chdir('/home/michael/Dropbox/written evaluations' )
    os.mkdir(lastname + ", " + firstname)
    os.chdir(lastname + ", " + firstname)

def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)

def create_ReferralLetter(PHYS):
    document = Document(r'/home/michael/Dropbox/psychpile/ReferralLetterTemplateBHP.docx' )
    for paragraph in document.paragraphs:
        find_replace("PTFN", firstname, paragraph)
        find_replace("PTLN", lastname, paragraph)
        find_replace("DOBI", DOB, paragraph)
        find_replace("DOI", DOI, paragraph)
        find_replace("DOT", DOT, paragraph)
        find_replace("DOR", DOR, paragraph)
        find_replace("PHYS", PHYS, paragraph)
        find_replace("GD", GD, paragraph)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
    for paragraph in document.paragraphs:
        paragraph.style = document.styles['Normal']

    if sys.platform == 'win32':
        os.chdir('C:\\Users\My PC\Dropbox\written evaluations' )
    else: os.chdir('/home/michael/Dropbox/written evaluations' )

    os.chdir(lastname + ", " + firstname)

    save_filename = "%s physicianletter.docx" % PHYS

    document.save(save_filename)

def create_faxsheet(PHYS):
    document = Document(r'/home/michael/Dropbox/psychpile/BHP faxsheet.docx' )
    for paragraph in document.paragraphs:
        find_replace("PTFN", firstname, paragraph)
        find_replace("PTLN", lastname, paragraph)
        find_replace("DOBI", DOB, paragraph)
        find_replace("DOI", DOI, paragraph)
        find_replace("DOT", DOT, paragraph)
        find_replace("DOR", DOR, paragraph)
        find_replace("PHYS", PHYS, paragraph)
        find_replace("GD", GD, paragraph)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    for paragraph in document.paragraphs:
        paragraph.style = document.styles['Normal']

    save_filename = "%s physicianfaxsheet.docx" % PHYS

    document.save(save_filename)

def create_eval():
    document = Document(r'/home/michael/Dropbox/psychpile/BHP letterhead.docx')
    for paragraph in document.paragraphs:
        find_replace("PTFN", firstname, paragraph)
        find_replace("PTLN", lastname, paragraph)
        find_replace("DOBI", DOB, paragraph)
        find_replace("DOI", DOI, paragraph)
        find_replace("DOT", DOT, paragraph)
        find_replace("DOR", DOR, paragraph)
        find_replace("PHYS", PHYS, paragraph)
        find_replace("GD1", GD1, paragraph)
        find_replace("GD2", GD2, paragraph)
        find_replace("GD3", GD3, paragraph)
        find_replace("GD", GD, paragraph)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    for paragraph in document.paragraphs:
        paragraph.style = document.styles['Normal']


    save_filename = "%s Full Evaluation.docx" % lastname
    document.save(save_filename)

    #now to edit the header

    WORKING_DIR = os.getcwd()
    TEMP_DOCX = os.path.join(WORKING_DIR, "%s Full Evaluation.docx" % lastname)
    TEMP_ZIP = os.path.join(WORKING_DIR, "template.zip")
    TEMP_FOLDER = os.path.join(WORKING_DIR, "template")

    # remove old zip file or folder template
    if os.path.exists(TEMP_ZIP):
        os.remove(TEMP_ZIP)
    if os.path.exists(TEMP_FOLDER):
        shutil.rmtree(TEMP_FOLDER)

    # reformat template.docx's extension
    os.rename(TEMP_DOCX, TEMP_ZIP)

    # unzip file zip to specific folder

    with zipfile.ZipFile(TEMP_ZIP, 'r') as z:
        z.extractall(TEMP_FOLDER)

    # change header xml file
    header_xml = os.path.join(TEMP_FOLDER, "word", "header1.xml")
    xmlstring = open(header_xml, 'r', encoding='utf-8').read()
    xmlstring = xmlstring.replace("PTLN", lastname)
    xmlstring = xmlstring.replace("PTFN", firstname)
    with open(header_xml, "wb") as f:
        f.write(xmlstring.encode("UTF-8"))

    # zip temp folder to zip file
    os.remove(TEMP_ZIP)
    shutil.make_archive(TEMP_ZIP.replace(".zip", ""), 'zip', TEMP_FOLDER)

    # rename zip file to docx
    os.rename(TEMP_ZIP, TEMP_DOCX)
    shutil.rmtree(TEMP_FOLDER)


layout = [
    [sg.Text('Lets Compile this PsychFile. Woooo.',justification='center',size=(70,1))],
    [sg.Text('Mr./ Ms/ Dr.', size =(20, 1)), sg.InputText(key='input_GD', do_not_clear=False)],
    [sg.Text('him or her', size =(20, 1)), sg.InputText(key='input_GD1', do_not_clear=False)],
    [sg.Text('he or she', size =(20, 1)), sg.InputText(key='input_GD2', do_not_clear=False)],
    [sg.Text('his or her', size =(20, 1)), sg.InputText(key='input_GD3', do_not_clear=False)],
    [sg.Text('Patient First Name', size =(20, 1)), sg.InputText(key='input_firstname', do_not_clear=False)],
    [sg.Text('Patient Last Name', size =(20, 1)), sg.InputText(key='input_lastname', do_not_clear=False)],
    [sg.Text('Date of Birth', size =(20, 1)), sg.InputText(key='input_DOB', do_not_clear=False)],
    [sg.Text('Clinical Interview Date', size =(20, 1)), sg.InputText(key='input_DOI', do_not_clear=False)],
    [sg.Text('Date of Testing', size =(20, 1)), sg.InputText(key='input_DOT', do_not_clear=False)],
    [sg.Text('Report Feedback Date', size =(20, 1)), sg.InputText(key='input_DOR', do_not_clear=False)],
    [sg.Text('Referrer', size =(20, 1)), sg.InputText(key='input_PHYS', do_not_clear=False)],
    [sg.Text('Referrer 2', size =(20, 1)), sg.InputText(key='input_PHYS_2', do_not_clear=False)],
  #  [sg.T("")], [sg.Text("Choose a  background file: "), sg.Input(), sg.FileBrowse(key="input_BKGRD-KEY-")],
   # [sg.T("")], [sg.Text("Choose a  testing file: "), sg.Input(), sg.FileBrowse(key="input_TESTING-KEY-")],
    [sg.Button("Submit")]]



window = sg.Window('PsychPile 1.2.3', layout)

while True:
    event, values = window.read()

    if event == sg.WIN_CLOSED:
        break

    if event is None or event == 'Cancel':
        break

    if event == 'Submit':

        GD = (values['input_GD'])
        GD1 = (values['input_GD1'])
        GD2 = (values['input_GD2'])
        GD3 = (values['input_GD3'])
        firstname = (values['input_firstname'])
        lastname = (values['input_lastname'])
        DOB = (values['input_DOB'])
        DOI = (values['input_DOI'])
        DOT = (values['input_DOT'])
        DOR = (values['input_DOR'])
        PHYS = (values['input_PHYS'])
        PHYS2 = (values['input_PHYS_2'])

#now make things
        create_folder()
        create_eval()
        create_faxsheet(PHYS)
        create_ReferralLetter(PHYS)
        if PHYS2 == None:
            break
            create_faxsheet(PHYS2)
            create_ReferralLetter(PHYS2)

        sg.popup('Complete. Please Close the Program Before Accessing Files')  # Shows OK button
window.close()
exit()
